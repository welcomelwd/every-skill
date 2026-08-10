#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Location: ./mcp-servers/python/url_to_markdown_server/src/url_to_markdown_server/server_fastmcp.py
Copyright 2025
SPDX-License-Identifier: Apache-2.0
Authors: Mihai Criveti

URL-to-Markdown FastMCP Server

A modern FastMCP implementation of the URL-to-Markdown converter with comprehensive
HTML, PDF, and document conversion capabilities.
"""

import asyncio
import contextvars
import logging
import mimetypes
import os
import re
import sys
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse
from uuid import uuid4

import httpcore
import httpx
from fastmcp import FastMCP
from httpcore._backends.auto import AutoBackend
from pydantic import Field

from .ssrf import SsrfBlockedError, _env_bool, validate_url

# Per-task pin of the validated destination IP for the in-flight fetch, read by
# _PinnedNetworkBackend.connect_tcp() below. A contextvar (not a converter
# attribute) is required because concurrent batch_convert() fetches share one
# converter/session but run as independent asyncio tasks, each needing its own
# pinned target.
_pinned_ip: contextvars.ContextVar[str | None] = contextvars.ContextVar("_pinned_ip", default=None)


class _PinnedNetworkBackend(httpcore.AsyncNetworkBackend):
    """TCP backend that connects to a per-request pinned IP instead of re-resolving DNS.

    The connection pool's request URL - and therefore its origin, connection-pool
    cache key, default TLS SNI, ``Host`` header, and cookie-jar scoping - stays the
    caller's real hostname; only the raw socket destination is substituted here via
    ``_pinned_ip``. Without this indirection, pinning by embedding the resolved IP
    directly in the request URL (the pre-fix approach) makes the IP itself the
    connection's origin: two different hostnames that validate_url() resolves to
    the *same* IP then collapse onto the same pooled origin, so a TLS connection
    handshaked and certificate-validated for the first hostname can be silently
    reused - via HTTPX/httpcore's own keep-alive pooling - to serve a request
    actually addressed to the second hostname, whose certificate was never
    checked. Keeping the origin as the hostname means HTTPX's per-origin pooling
    naturally stays scoped per hostname, exactly like an unpinned client; only the
    DNS step is substituted with the address ``ssrf.py`` already validated,
    closing the DNS-rebinding window between validation and connect.
    """

    def __init__(self) -> None:
        """Wrap a fresh AutoBackend for the real connect/TLS/sleep work."""
        self._backend: httpcore.AsyncNetworkBackend = AutoBackend()

    async def connect_tcp(
        self,
        host: str,
        port: int,
        timeout: float | None = None,
        local_address: str | None = None,
        socket_options: Any = None,
    ) -> httpcore.AsyncNetworkStream:
        """Connect to the pinned IP for the in-flight request, if one is set."""
        pinned = _pinned_ip.get()
        target_host = pinned if pinned is not None else host
        return await self._backend.connect_tcp(
            target_host,
            port,
            timeout=timeout,
            local_address=local_address,
            socket_options=socket_options,
        )

    async def connect_unix_socket(
        self,
        path: str,
        timeout: float | None = None,
        socket_options: Any = None,
    ) -> httpcore.AsyncNetworkStream:
        """Delegate unix-socket connects unchanged (not used by this server)."""
        return await self._backend.connect_unix_socket(
            path, timeout=timeout, socket_options=socket_options
        )

    async def sleep(self, seconds: float) -> None:
        """Delegate retry backoff sleeps to plain asyncio."""
        await asyncio.sleep(seconds)


def _environment_proxy_mounts() -> dict[str, str | None] | None:
    """Return HTTPX's own environment-proxy mount map, or ``None`` if unavailable.

    ``httpx._utils.get_environment_proxies()`` is the exact function ``httpx.AsyncClient``
    uses internally to turn ``HTTP_PROXY``/``HTTPS_PROXY``/``ALL_PROXY``/``NO_PROXY`` into a
    ``{url_pattern: proxy_url_or_None}`` map - a ``None`` value means "use the default
    transport for this pattern instead of a proxy" (this is how ``NO_PROXY`` entries are
    represented). Reusing it, rather than re-parsing the env vars ourselves, guarantees the
    per-destination proxy decision in ``get_session()`` matches HTTPX's own request routing
    exactly, including ``NO_PROXY``'s CIDR/wildcard/``localhost`` handling.

    This is a private API. If a future HTTPX release removes it, fail closed: ignore the
    proxy environment variables and keep IP pinning on every destination, rather than
    silently reintroducing the DNS-rebinding gap this guards against.
    """
    try:
        from httpx._utils import get_environment_proxies  # pylint: disable=import-outside-toplevel

        return get_environment_proxies()
    except Exception:  # pylint: disable=broad-except
        logger.error(
            "Could not read environment proxy configuration via httpx; ignoring "
            "HTTP_PROXY/HTTPS_PROXY/ALL_PROXY/NO_PROXY and keeping IP pinning for all destinations."
        )
        return None


# Configure logging to stderr to avoid MCP protocol interference
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.StreamHandler(sys.stderr)],
)
logger = logging.getLogger(__name__)

# Configuration constants
DEFAULT_TIMEOUT = int(os.getenv("MARKDOWN_DEFAULT_TIMEOUT", "30"))
MAX_TIMEOUT = int(os.getenv("MARKDOWN_MAX_TIMEOUT", "120"))
MAX_CONTENT_SIZE = int(os.getenv("MARKDOWN_MAX_CONTENT_SIZE", "50971520"))  # 50MB
MAX_REDIRECT_HOPS = int(os.getenv("MARKDOWN_MAX_REDIRECT_HOPS", "10"))
DEFAULT_USER_AGENT = os.getenv("MARKDOWN_USER_AGENT", "URL-to-Markdown-MCP-Server/2.0")

# Create FastMCP server instance
mcp = FastMCP(name="url-to-markdown-server", version="2.0.0")


class UrlToMarkdownConverter:
    """Main converter class for URL-to-Markdown operations."""

    def __init__(self):
        """Initialize the converter."""
        self.session = None
        # The pinned direct transport built by get_session(), used to tell whether a
        # given request actually went out pinned or via an unpinned proxy mount - see
        # fetch_url_content(). None until the first get_session() call.
        self._pinned_transport: httpx.AsyncHTTPTransport | None = None
        self.html_engines = self._check_html_engines()
        self.document_converters = self._check_document_converters()

    def _check_html_engines(self) -> dict[str, bool]:
        """Check availability of HTML-to-Markdown engines."""
        engines = {}

        try:
            import html2text

            engines["html2text"] = True
        except ImportError:
            engines["html2text"] = False

        try:
            import markdownify

            engines["markdownify"] = True
        except ImportError:
            engines["markdownify"] = False

        try:
            from bs4 import BeautifulSoup

            engines["beautifulsoup"] = True
        except ImportError:
            engines["beautifulsoup"] = False

        try:
            from readability import Document

            engines["readability"] = True
        except ImportError:
            engines["readability"] = False

        return engines

    def _check_document_converters(self) -> dict[str, bool]:
        """Check availability of document converters."""
        converters = {}

        try:
            import pypandoc

            converters["pandoc"] = True
        except ImportError:
            converters["pandoc"] = False

        try:
            import fitz  # PyMuPDF

            converters["pymupdf"] = True
        except ImportError:
            converters["pymupdf"] = False

        try:
            from docx import Document

            converters["python_docx"] = True
        except ImportError:
            converters["python_docx"] = False

        try:
            import openpyxl

            converters["openpyxl"] = True
        except ImportError:
            converters["openpyxl"] = False

        return converters

    async def get_session(self) -> httpx.AsyncClient:
        """Get or create HTTP session.

        The client's default transport is always the pinned ``AsyncHTTPTransport``
        (network backend swapped for ``_PinnedNetworkBackend``), so any destination not
        otherwise routed - including every ``NO_PROXY``-excluded host - gets its raw TCP
        connect pinned to the IP ``ssrf.py`` already validated.

        When ``HTTP_PROXY``/``HTTPS_PROXY``/``ALL_PROXY`` are configured, the matching
        destinations are additionally mounted onto dedicated, unpinned proxy transports
        (mirroring HTTPX's own environment-proxy routing - see
        ``_environment_proxy_mounts()``): the socket for those requests would be opened
        to the proxy, which performs its own name resolution for the real destination,
        so the validated address cannot be enforced end-to-end on that specific path.
        The mount is still built unconditionally - so the session's routing accurately
        reflects the operator's proxy configuration, and ``NO_PROXY``/unmatched
        destinations are unaffected either way - but ``fetch_url_content()`` refuses to
        actually send a request over it unless ``MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING`` is
        set: silently falling back to a direct connection instead would bypass whatever
        egress policy (audit logging, DLP, destination allowlisting) the operator
        configured the proxy to enforce, which is its own kind of control bypass distinct
        from the DNS-rebinding risk pinning addresses. ``NO_PROXY`` entries are mounted
        back onto the pinned default transport instead of being left to fall through to
        a broader proxy mount, so a proxy configured for some destinations never weakens
        pinning for destinations that bypass it.
        """
        if self.session is None or self.session.is_closed:
            client_kwargs: dict[str, Any] = {
                "headers": {
                    "User-Agent": DEFAULT_USER_AGENT,
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,*/*;q=0.8",
                    "Accept-Language": "en-US,en;q=0.5",
                    "Accept-Encoding": "gzip, deflate",
                    "Connection": "keep-alive",
                    "Upgrade-Insecure-Requests": "1",
                },
                "timeout": httpx.Timeout(DEFAULT_TIMEOUT),
                # Redirects are followed manually in fetch_url_content so every hop can be
                # re-validated against the SSRF guard - a public host
                # that 302s to an internal address must not bypass validation on the second hop.
                "follow_redirects": False,
            }

            transport = httpx.AsyncHTTPTransport()
            # HTTPX's AsyncHTTPTransport has no public constructor option to inject a
            # custom network backend, so swapping the pool's backend after construction
            # is the only way to pin the raw TCP destination (see _PinnedNetworkBackend)
            # without reimplementing the transport's SSL/UDS setup ourselves.
            transport._pool._network_backend = (
                _PinnedNetworkBackend()
            )  # pylint: disable=protected-access
            self._pinned_transport = transport

            mounts: dict[str, httpx.AsyncBaseTransport | None] = {}
            proxied_patterns = []
            for pattern, proxy_url in (_environment_proxy_mounts() or {}).items():
                if proxy_url is None:
                    # A NO_PROXY-derived entry: pin this pattern explicitly to the
                    # default transport instead of letting it fall through to a
                    # broader proxy mount (e.g. "all://") that would otherwise match it.
                    mounts[pattern] = None
                else:
                    proxied_patterns.append(pattern)
                    # Unpinned: the proxy resolves the real destination itself, so
                    # connect-time IP pinning cannot be enforced end-to-end here.
                    # fetch_url_content() refuses to actually use this mount unless
                    # MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING is set.
                    mounts[pattern] = httpx.AsyncHTTPTransport(proxy=proxy_url)

            if proxied_patterns:
                logger.warning(
                    f"Egress proxy configured for {', '.join(sorted(proxied_patterns))}: "
                    "matching requests are refused before any network I/O unless "
                    "MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING is set, since connect-time IP "
                    "pinning cannot be enforced on that path (the proxy resolves the "
                    "destination itself). With that opt-in, matching requests are "
                    "routed through the proxy instead, unpinned. NO_PROXY-excluded and "
                    "otherwise-unmatched destinations keep IP pinning either way."
                )

            self.session = httpx.AsyncClient(
                transport=transport, mounts=mounts or None, **client_kwargs
            )
        return self.session

    async def fetch_url_content(self, url: str, timeout: int = DEFAULT_TIMEOUT) -> dict[str, Any]:
        """Fetch content from URL with comprehensive error handling.

        Every hop of a redirect chain is independently re-validated against the SSRF
        guard. The request is made to the original hostname (so HTTPX's connection
        pool, TLS SNI, ``Host`` header, and cookie jar all stay keyed by hostname as
        usual) while the raw TCP connect is pinned to the resolved IP via
        ``_pinned_ip``/``_PinnedNetworkBackend``, closing the DNS-rebinding window
        between validation and connect without letting two hostnames that share a
        resolved IP collapse onto the same pooled connection. The response body is
        streamed and the transfer is aborted as soon as it exceeds
        ``MAX_CONTENT_SIZE`` rather than buffered in full first.

        A destination that matches a configured egress proxy is refused - before any
        connection to either the proxy or the destination - unless
        ``MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING`` is set: routing it through the proxy
        would leave its connect unpinned (see ``get_session()``), and falling back to a
        direct connection instead would silently bypass whatever egress policy the
        operator configured the proxy to enforce. ``NO_PROXY``-excluded and otherwise-
        unmatched destinations always get their connect pinned and are unaffected by
        this flag. Validation and per-hop re-validation are unchanged either way.
        """
        session = await self.get_session()
        current_url = url
        allow_unsafe_proxy_pinning = _env_bool("MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING", False)

        try:
            for _ in range(MAX_REDIRECT_HOPS + 1):
                try:
                    target = await validate_url(current_url, timeout=timeout)
                except SsrfBlockedError as e:
                    logger.warning(f"Blocked URL (SSRF guard): {current_url}")
                    return {"success": False, "error": str(e)}

                pinned = (
                    session._transport_for_url(  # pylint: disable=protected-access
                        httpx.URL(target.validated_url)
                    )
                    is self._pinned_transport
                )
                if not pinned and not allow_unsafe_proxy_pinning:
                    logger.warning(
                        f"Blocked URL (egress proxy without MARKDOWN_ALLOW_UNSAFE_PROXY_PINNING): "
                        f"{target.hostname}"
                    )
                    return {"success": False, "error": "URL is not allowed"}

                mode = "pinned" if pinned else "via proxy, unpinned"
                logger.info(f"Fetching URL ({mode}): {target.hostname} -> {target.resolved_ip}")

                pin_token = _pinned_ip.set(target.resolved_ip)
                try:
                    async with session.stream(
                        "GET",
                        target.validated_url,
                        timeout=timeout,
                    ) as response:
                        if response.status_code in (301, 302, 303, 307, 308):
                            location = response.headers.get("location")
                            if not location:
                                return {
                                    "success": False,
                                    "error": "Redirect missing Location header",
                                }
                            current_url = urljoin(target.validated_url, location)
                            continue

                        response.raise_for_status()

                        # Check content size
                        content_length = response.headers.get("content-length")
                        if content_length and int(content_length) > MAX_CONTENT_SIZE:
                            return {
                                "success": False,
                                "error": (
                                    f"Content too large: {content_length} bytes "
                                    f"(max: {MAX_CONTENT_SIZE})"
                                ),
                            }

                        chunks: list[bytes] = []
                        total = 0
                        async for chunk in response.aiter_bytes():
                            total += len(chunk)
                            if total > MAX_CONTENT_SIZE:
                                return {
                                    "success": False,
                                    "error": (
                                        f"Content too large: exceeded {MAX_CONTENT_SIZE} "
                                        "bytes while streaming"
                                    ),
                                }
                            chunks.append(chunk)
                        content = b"".join(chunks)

                        # Determine content type
                        content_type = response.headers.get("content-type", "").lower()
                        detected_type = self._detect_content_type(
                            content, content_type, target.validated_url
                        )

                        return {
                            "success": True,
                            "content": content,
                            "content_type": detected_type,
                            "original_content_type": content_type,
                            "url": target.validated_url,  # Final URL after redirects
                            "status_code": response.status_code,
                            "headers": dict(response.headers),
                            "size": len(content),
                        }
                except httpx.TimeoutException:
                    return {"success": False, "error": f"Request timeout after {timeout} seconds"}
                except httpx.HTTPStatusError as e:
                    return {
                        "success": False,
                        "error": f"HTTP {e.response.status_code}: {e.response.reason_phrase}",
                    }
                finally:
                    # Never let one hop's pin leak into the next hop's connect - each
                    # redirect target is revalidated and gets its own pinned IP above.
                    _pinned_ip.reset(pin_token)

            return {"success": False, "error": f"Too many redirects (max {MAX_REDIRECT_HOPS})"}

        except Exception as e:
            logger.error(f"Error fetching URL {url}: {e}")
            return {"success": False, "error": str(e)}

    def _detect_content_type(self, content: bytes, content_type: str, url: str) -> str:
        """Detect actual content type from content, headers, and URL."""
        # Check file extension first
        url_path = urlparse(url).path.lower()

        if url_path.endswith((".pdf",)):
            return "application/pdf"
        elif url_path.endswith((".docx",)):
            return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif url_path.endswith((".txt", ".md", ".rst")):
            return "text/plain"

        # Check content-type header
        if "html" in content_type:
            return "text/html"
        elif "pdf" in content_type:
            return "application/pdf"
        elif "json" in content_type:
            return "application/json"

        # Check magic bytes
        if content.startswith(b"%PDF"):
            return "application/pdf"
        elif content.startswith(b"PK"):  # ZIP-based formats
            if b"word/" in content[:1024]:
                return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif content.startswith((b"<html", b"<!DOCTYPE", b"<!doctype")):
            return "text/html"

        return content_type or "application/octet-stream"

    async def convert_html_to_markdown(
        self,
        html_content: str,
        base_url: str = "",
        engine: str = "html2text",
        include_images: bool = True,
        include_links: bool = True,
    ) -> dict[str, Any]:
        """Convert HTML content to markdown using specified engine."""
        try:
            if engine == "html2text" and self.html_engines.get("html2text"):
                return await self._convert_with_html2text(
                    html_content, base_url, include_images, include_links
                )
            elif engine == "markdownify" and self.html_engines.get("markdownify"):
                return await self._convert_with_markdownify(
                    html_content, include_images, include_links
                )
            elif engine == "beautifulsoup" and self.html_engines.get("beautifulsoup"):
                return await self._convert_with_beautifulsoup(
                    html_content, base_url, include_images
                )
            elif engine == "readability" and self.html_engines.get("readability"):
                return await self._convert_with_readability(html_content, base_url)
            else:
                # Fallback to basic conversion
                return await self._convert_basic_html(html_content)

        except Exception as e:
            logger.error(f"Error converting HTML to markdown: {e}")
            return {"success": False, "error": f"Conversion failed: {str(e)}"}

    async def _convert_with_html2text(
        self, html_content: str, base_url: str, include_images: bool, include_links: bool
    ) -> dict[str, Any]:
        """Convert using html2text library."""
        import html2text

        converter = html2text.HTML2Text()
        converter.ignore_links = not include_links
        converter.ignore_images = not include_images
        converter.body_width = 0  # No line wrapping
        converter.protect_links = True
        converter.wrap_links = False

        if base_url:
            converter.baseurl = base_url

        markdown = converter.handle(html_content)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "html2text",
            "length": len(markdown),
        }

    async def _convert_with_markdownify(
        self, html_content: str, include_images: bool, include_links: bool
    ) -> dict[str, Any]:
        """Convert using markdownify library."""
        import markdownify

        # Configure conversion options
        options = {
            "heading_style": "ATX",  # Use # for headings
            "bullets": "-",  # Use - for lists
            "escape_misc": False,
        }

        if not include_links:
            options["convert"] = [
                "p",
                "div",
                "span",
                "h1",
                "h2",
                "h3",
                "h4",
                "h5",
                "h6",
                "ul",
                "ol",
                "li",
            ]

        if not include_images:
            if "convert" in options:
                pass  # img already excluded
            else:
                options["strip"] = ["img"]

        markdown = markdownify.markdownify(html_content, **options)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "markdownify",
            "length": len(markdown),
        }

    async def _convert_with_beautifulsoup(
        self, html_content: str, base_url: str, include_images: bool
    ) -> dict[str, Any]:
        """Convert using BeautifulSoup for parsing + custom markdown generation."""
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html_content, "html.parser")

        # Extract main content
        main_content = self._extract_main_content(soup)

        # Convert to markdown
        markdown = self._soup_to_markdown(main_content, base_url, include_images)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "beautifulsoup",
            "length": len(markdown),
        }

    async def _convert_with_readability(self, html_content: str, base_url: str) -> dict[str, Any]:
        """Convert using readability for content extraction."""
        from readability import Document

        doc = Document(html_content)
        title = doc.title()
        content = doc.summary()

        # Convert extracted content to markdown
        if self.html_engines.get("html2text"):
            import html2text

            converter = html2text.HTML2Text()
            converter.body_width = 0
            if base_url:
                converter.baseurl = base_url
            markdown = converter.handle(content)
        else:
            # Basic conversion
            markdown = self._html_to_markdown_basic(content)

        # Add title if available
        if title:
            markdown = f"# {title}\n\n{markdown}"

        return {
            "success": True,
            "markdown": markdown,
            "engine": "readability",
            "title": title,
            "length": len(markdown),
        }

    async def _convert_basic_html(self, html_content: str) -> dict[str, Any]:
        """Basic HTML to markdown conversion without external libraries."""
        markdown = self._html_to_markdown_basic(html_content)

        return {
            "success": True,
            "markdown": markdown,
            "engine": "basic",
            "length": len(markdown),
            "note": "Basic conversion - install html2text or markdownify for better results",
        }

    def _html_to_markdown_basic(self, html_content: str) -> str:
        """Basic HTML to markdown conversion."""
        # Remove script and style tags
        html_content = re.sub(
            r"<script[^>]*>.*?</script>", "", html_content, flags=re.DOTALL | re.IGNORECASE
        )
        html_content = re.sub(
            r"<style[^>]*>.*?</style>", "", html_content, flags=re.DOTALL | re.IGNORECASE
        )

        # Convert headings
        for i in range(1, 7):
            html_content = re.sub(
                f"<h{i}[^>]*>(.*?)</h{i}>",
                f"{'#' * i} \\1\n\n",
                html_content,
                flags=re.DOTALL | re.IGNORECASE,
            )

        # Convert paragraphs
        html_content = re.sub(
            r"<p[^>]*>(.*?)</p>", r"\1\n\n", html_content, flags=re.DOTALL | re.IGNORECASE
        )

        # Convert line breaks
        html_content = re.sub(r"<br[^>]*/?>", "\n", html_content, flags=re.IGNORECASE)

        # Convert links
        html_content = re.sub(
            r'<a[^>]*href=["\']([^"\']+)["\'][^>]*>(.*?)</a>',
            r"[\2](\1)",
            html_content,
            flags=re.DOTALL | re.IGNORECASE,
        )

        # Convert bold and italic
        html_content = re.sub(
            r"<(strong|b)[^>]*>(.*?)</\1>", r"**\2**", html_content, flags=re.DOTALL | re.IGNORECASE
        )
        html_content = re.sub(
            r"<(em|i)[^>]*>(.*?)</\1>", r"*\2*", html_content, flags=re.DOTALL | re.IGNORECASE
        )

        # Convert lists
        html_content = re.sub(
            r"<li[^>]*>(.*?)</li>", r"- \1\n", html_content, flags=re.DOTALL | re.IGNORECASE
        )
        html_content = re.sub(r"<[uo]l[^>]*>", "\n", html_content, flags=re.IGNORECASE)
        html_content = re.sub(r"</[uo]l>", "\n", html_content, flags=re.IGNORECASE)

        # Remove remaining HTML tags
        html_content = re.sub(r"<[^>]+>", "", html_content)

        # Clean up whitespace
        html_content = re.sub(r"\n\s*\n\s*\n", "\n\n", html_content)
        html_content = re.sub(r"^\s+|\s+$", "", html_content, flags=re.MULTILINE)

        return html_content.strip()

    def _extract_main_content(self, soup):
        """Extract main content from BeautifulSoup object."""
        # Try to find main content areas
        main_selectors = [
            "main",
            "article",
            '[role="main"]',
            ".content",
            ".main-content",
            ".post-content",
            "#content",
            "#main-content",
            "#post-content",
        ]

        for selector in main_selectors:
            main_element = soup.select_one(selector)
            if main_element:
                return main_element

        # Fallback to body
        body = soup.find("body")
        if body:
            # Remove navigation, sidebar, footer elements
            for element in body.find_all(["nav", "aside", "footer", "header"]):
                element.decompose()

            # Remove elements with common nav/sidebar classes
            for element in body.find_all(
                class_=re.compile(r"(nav|sidebar|footer|header|menu)", re.I)
            ):
                element.decompose()

            return body

        return soup

    def _soup_to_markdown(self, element, base_url: str = "", include_images: bool = True) -> str:
        """Convert BeautifulSoup element to markdown."""
        markdown_parts = []

        for child in element.children:
            if hasattr(child, "name"):
                if child.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    level = int(child.name[1])
                    text = child.get_text().strip()
                    markdown_parts.append(f"{'#' * level} {text}\n")
                elif child.name == "p":
                    text = child.get_text().strip()
                    if text:
                        markdown_parts.append(f"{text}\n")
                elif child.name == "a":
                    href = child.get("href", "")
                    text = child.get_text().strip()
                    if href and text:
                        if base_url and not href.startswith(("http", "https")):
                            href = urljoin(base_url, href)
                        markdown_parts.append(f"[{text}]({href})")
                elif child.name == "img" and include_images:
                    src = child.get("src", "")
                    alt = child.get("alt", "Image")
                    if src:
                        if base_url and not src.startswith(("http", "https")):
                            src = urljoin(base_url, src)
                        markdown_parts.append(f"![{alt}]({src})")
                elif child.name in ["strong", "b"]:
                    text = child.get_text().strip()
                    markdown_parts.append(f"**{text}**")
                elif child.name in ["em", "i"]:
                    text = child.get_text().strip()
                    markdown_parts.append(f"*{text}*")
                elif child.name == "li":
                    text = child.get_text().strip()
                    markdown_parts.append(f"- {text}\n")
                elif child.name == "code":
                    text = child.get_text()
                    markdown_parts.append(f"`{text}`")
                elif child.name == "pre":
                    text = child.get_text()
                    markdown_parts.append(f"```\n{text}\n```\n")
                else:
                    # Recursively process other elements
                    nested_markdown = self._soup_to_markdown(child, base_url, include_images)
                    if nested_markdown.strip():
                        markdown_parts.append(nested_markdown)
            else:
                # Text node
                text = str(child).strip()
                if text:
                    markdown_parts.append(text)

        return " ".join(markdown_parts)

    async def convert_document_to_markdown(
        self, content: bytes, content_type: str
    ) -> dict[str, Any]:
        """Convert document formats to markdown."""
        try:
            if content_type == "application/pdf":
                return await self._convert_pdf_to_markdown(content)
            elif "wordprocessingml" in content_type:  # DOCX
                return await self._convert_docx_to_markdown(content)
            elif content_type.startswith("text/"):
                return await self._convert_text_to_markdown(content)
            else:
                return {"success": False, "error": f"Unsupported content type: {content_type}"}

        except Exception as e:
            logger.error(f"Error converting document: {e}")
            return {"success": False, "error": f"Document conversion failed: {str(e)}"}

    async def _convert_pdf_to_markdown(self, pdf_content: bytes) -> dict[str, Any]:
        """Convert PDF to markdown."""
        if not self.document_converters.get("pymupdf"):
            return {"success": False, "error": "PyMuPDF not available for PDF conversion"}

        try:
            import fitz

            # Open PDF from bytes
            doc = fitz.open(stream=pdf_content, filetype="pdf")

            markdown_parts = []

            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()

                if text.strip():
                    markdown_parts.append(f"## Page {page_num + 1}\n\n{text}\n")

            doc.close()

            markdown = "\n".join(markdown_parts)

            return {
                "success": True,
                "markdown": markdown,
                "engine": "pymupdf",
                "pages": len(doc),
                "length": len(markdown),
            }

        except Exception as e:
            return {"success": False, "error": f"PDF conversion error: {str(e)}"}

    async def _convert_docx_to_markdown(self, docx_content: bytes) -> dict[str, Any]:
        """Convert DOCX to markdown."""
        if not self.document_converters.get("python_docx"):
            return {"success": False, "error": "python-docx not available for DOCX conversion"}

        try:
            from io import BytesIO

            from docx import Document

            doc = Document(BytesIO(docx_content))
            markdown_parts = []

            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    # Check if it's a heading based on style
                    if paragraph.style.name.startswith("Heading"):
                        level = int(paragraph.style.name.split()[-1])
                        markdown_parts.append(f"{'#' * level} {text}\n")
                    else:
                        markdown_parts.append(f"{text}\n")

            markdown = "\n".join(markdown_parts)

            return {
                "success": True,
                "markdown": markdown,
                "engine": "python_docx",
                "paragraphs": len(doc.paragraphs),
                "length": len(markdown),
            }

        except Exception as e:
            return {"success": False, "error": f"DOCX conversion error: {str(e)}"}

    async def _convert_text_to_markdown(self, text_content: bytes) -> dict[str, Any]:
        """Convert plain text to markdown."""
        try:
            text = text_content.decode("utf-8", errors="replace")

            # For plain text, just return as-is with minimal formatting
            markdown = text

            return {
                "success": True,
                "markdown": markdown,
                "engine": "text",
                "length": len(markdown),
            }

        except Exception as e:
            return {"success": False, "error": f"Text conversion error: {str(e)}"}

    def get_capabilities(self) -> dict[str, Any]:
        """Get converter capabilities and available engines."""
        return {
            "html_engines": self.html_engines,
            "document_converters": self.document_converters,
            "supported_formats": {
                "web": ["text/html", "application/xhtml+xml"],
                "documents": ["application/pdf"],
                "office": [
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",  # DOCX
                ],
                "text": ["text/plain", "text/markdown", "application/json"],
            },
            "features": [
                "Multi-engine HTML conversion",
                "PDF text extraction",
                "Office document conversion",
                "Content cleaning and optimization",
                "Image handling",
                "Link preservation",
                "Batch processing",
                "Metadata extraction",
            ],
            "configuration": {
                "default_timeout": DEFAULT_TIMEOUT,
                "max_timeout": MAX_TIMEOUT,
                "max_content_size": MAX_CONTENT_SIZE,
                "max_redirect_hops": MAX_REDIRECT_HOPS,
                "user_agent": DEFAULT_USER_AGENT,
            },
        }

    def clean_markdown(self, markdown: str) -> str:
        """Clean and optimize markdown content."""
        # Remove excessive whitespace
        markdown = re.sub(r"\n\s*\n\s*\n+", "\n\n", markdown)

        # Fix heading spacing
        markdown = re.sub(r"(#+\s+.+)\n+([^#\n])", r"\1\n\n\2", markdown)

        # Clean up list formatting
        markdown = re.sub(r"\n+(-\s+)", r"\n\1", markdown)

        # Remove empty links
        markdown = re.sub(r"\[\s*\]\([^)]*\)", "", markdown)

        # Clean up extra spaces
        markdown = re.sub(r" +", " ", markdown)

        # Trim
        return markdown.strip()


# Initialize global converter
converter = UrlToMarkdownConverter()


# Tool definitions using FastMCP
@mcp.tool(description="Convert URL content to markdown format with multiple engines and options")
async def convert_url(
    url: str = Field(..., description="URL to retrieve and convert to markdown"),
    timeout: int = Field(DEFAULT_TIMEOUT, le=MAX_TIMEOUT, description="Request timeout in seconds"),
    include_images: bool = Field(True, description="Include images in markdown"),
    include_links: bool = Field(True, description="Preserve links in markdown"),
    clean_content: bool = Field(True, description="Clean and optimize content"),
    extraction_method: str = Field(
        "auto", pattern="^(auto|readability|raw)$", description="HTML extraction method"
    ),
    markdown_engine: str = Field(
        "html2text",
        pattern="^(html2text|markdownify|beautifulsoup|basic)$",
        description="Markdown conversion engine",
    ),
) -> dict[str, Any]:
    """Convert a URL to markdown with comprehensive format support."""
    conversion_id = str(uuid4())
    logger.info(f"Converting URL to markdown, ID: {conversion_id}, URL: {url}")

    try:
        # Fetch content
        fetch_result = await converter.fetch_url_content(url, timeout)
        if not fetch_result["success"]:
            return {
                "success": False,
                "conversion_id": conversion_id,
                "error": fetch_result["error"],
            }

        content = fetch_result["content"]
        content_type = fetch_result["content_type"]
        final_url = fetch_result["url"]

        # Convert based on content type
        if content_type.startswith("text/html"):
            html_content = content.decode("utf-8", errors="replace")

            # Choose extraction method
            if extraction_method == "readability":
                result = await converter._convert_with_readability(html_content, final_url)
            elif extraction_method == "raw":
                result = await converter.convert_html_to_markdown(
                    html_content, final_url, markdown_engine, include_images, include_links
                )
            else:  # auto
                # Try readability first, fallback to specified engine
                if converter.html_engines.get("readability"):
                    result = await converter._convert_with_readability(html_content, final_url)
                else:
                    result = await converter.convert_html_to_markdown(
                        html_content, final_url, markdown_engine, include_images, include_links
                    )
        else:
            # Handle document formats
            result = await converter.convert_document_to_markdown(content, content_type)

        if not result["success"]:
            return {"success": False, "conversion_id": conversion_id, "error": result["error"]}

        markdown = result["markdown"]

        # Clean content if requested
        if clean_content:
            markdown = converter.clean_markdown(markdown)

        return {
            "success": True,
            "conversion_id": conversion_id,
            "url": final_url,
            "content_type": content_type,
            "markdown": markdown,
            "length": len(markdown),
            "engine": result.get("engine", "unknown"),
            "metadata": {
                "original_size": len(content),
                "compression_ratio": len(markdown) / len(content) if len(content) > 0 else 0,
            },
        }

    except Exception as e:
        logger.error(f"Error converting URL {url}: {e}")
        return {"success": False, "conversion_id": conversion_id, "error": str(e)}


@mcp.tool(description="Convert raw HTML or text content to markdown")
async def convert_content(
    content: str = Field(..., description="Raw content to convert to markdown"),
    content_type: str = Field("text/html", description="MIME type of the content"),
    base_url: str | None = Field(None, description="Base URL for resolving relative links"),
    include_images: bool = Field(True, description="Include images in markdown"),
    clean_content: bool = Field(True, description="Clean and optimize content"),
    markdown_engine: str = Field(
        "html2text",
        pattern="^(html2text|markdownify|beautifulsoup|basic)$",
        description="Markdown conversion engine",
    ),
) -> dict[str, Any]:
    """Convert raw content to markdown format."""
    try:
        if content_type.startswith("text/html"):
            result = await converter.convert_html_to_markdown(
                html_content=content,
                base_url=base_url or "",
                engine=markdown_engine,
                include_images=include_images,
            )
        else:
            result = await converter.convert_document_to_markdown(
                content=content.encode("utf-8"), content_type=content_type
            )

        if result["success"] and clean_content:
            result["markdown"] = converter.clean_markdown(result["markdown"])

        return result

    except Exception as e:
        logger.error(f"Error converting content: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool(description="Convert a local file to markdown format")
async def convert_file(
    file_path: str = Field(..., description="Path to local file to convert"),
    include_images: bool = Field(True, description="Include images in markdown"),
    clean_content: bool = Field(True, description="Clean and optimize content"),
) -> dict[str, Any]:
    """Convert a local file to markdown."""
    try:
        file_path_obj = Path(file_path)
        if not file_path_obj.exists():
            return {"success": False, "error": f"File not found: {file_path}"}

        content = file_path_obj.read_bytes()
        content_type = mimetypes.guess_type(str(file_path_obj))[0] or "application/octet-stream"

        result = await converter.convert_document_to_markdown(content, content_type)

        if result["success"] and clean_content:
            result["markdown"] = converter.clean_markdown(result["markdown"])

        result["file_path"] = str(file_path_obj)
        return result

    except Exception as e:
        logger.error(f"Error converting file {file_path}: {e}")
        return {"success": False, "error": str(e)}


@mcp.tool(description="Convert multiple URLs to markdown in parallel")
async def batch_convert(
    urls: list[str] = Field(
        ..., max_length=50, description="List of URLs to convert to markdown (max 50 per batch)"
    ),
    timeout: int = Field(DEFAULT_TIMEOUT, le=MAX_TIMEOUT, description="Request timeout per URL"),
    max_concurrent: int = Field(5, ge=1, le=10, description="Maximum concurrent requests"),
    include_images: bool = Field(False, description="Include images in markdown"),
    clean_content: bool = Field(True, description="Clean and optimize content"),
) -> dict[str, Any]:
    """Batch convert multiple URLs to markdown concurrently."""
    batch_id = str(uuid4())
    logger.info(f"Batch converting {len(urls)} URLs, ID: {batch_id}")

    semaphore = asyncio.Semaphore(max_concurrent)

    async def convert_single_url(url: str) -> dict[str, Any]:
        async with semaphore:
            return await convert_url(
                url=url,
                timeout=timeout,
                include_images=include_images,
                include_links=True,
                clean_content=clean_content,
            )

    try:
        # Process URLs concurrently
        tasks = [convert_single_url(url) for url in urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # Process results
        successful = 0
        failed = 0
        processed_results = []

        for i, result in enumerate(results):
            if isinstance(result, Exception):
                processed_results.append({"url": urls[i], "success": False, "error": str(result)})
                failed += 1
            else:
                processed_results.append(result)
                if result.get("success"):
                    successful += 1
                else:
                    failed += 1

        return {
            "success": True,
            "batch_id": batch_id,
            "total_urls": len(urls),
            "successful": successful,
            "failed": failed,
            "results": processed_results,
        }

    except Exception as e:
        logger.error(f"Error in batch conversion: {e}")
        return {"success": False, "batch_id": batch_id, "error": str(e)}


@mcp.tool(description="Get information about converter capabilities and available engines")
async def get_capabilities() -> dict[str, Any]:
    """Get converter capabilities and available engines."""
    return converter.get_capabilities()


def main():
    """Main server entry point."""
    import argparse

    parser = argparse.ArgumentParser(description="URL-to-Markdown FastMCP Server")
    parser.add_argument(
        "--transport",
        choices=["stdio", "http"],
        default="stdio",
        help="Transport mode (stdio or http)",
    )
    parser.add_argument("--host", default="0.0.0.0", help="HTTP host")
    parser.add_argument("--port", type=int, default=9016, help="HTTP port")

    args = parser.parse_args()

    if args.transport == "http":
        logger.info(f"Starting URL-to-Markdown FastMCP Server on HTTP at {args.host}:{args.port}")
        mcp.run(transport="http", host=args.host, port=args.port)
    else:
        logger.info("Starting URL-to-Markdown FastMCP Server on stdio")
        mcp.run()


if __name__ == "__main__":
    main()
